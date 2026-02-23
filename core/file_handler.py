"""Document import and export utilities.

Supports reading .md, .txt, .docx files and writing .md and .docx output.
"""

import logging
from pathlib import Path

import markdown2
from docx import Document
from markdownify import markdownify as md_from_html

logger = logging.getLogger(__name__)

SUPPORTED_READ_EXTENSIONS = {".md", ".txt", ".docx"}


class FileHandlerError(Exception):
    """Raised when a file cannot be read or written."""


# ------------------------------------------------------------------
# Reading
# ------------------------------------------------------------------

def read_file(path: Path) -> str:
    """Read a document and return its contents as plain text.

    Dispatches to the appropriate parser based on file extension.

    Args:
        path: Path to the file to read.

    Returns:
        The document's text content as a UTF-8 string.

    Raises:
        FileHandlerError: If the file type is unsupported or cannot be read.
    """
    path = Path(path)
    if not path.exists():
        raise FileHandlerError(f"File not found: {path}")

    ext = path.suffix.lower()
    if ext in (".md", ".txt"):
        return _read_text(path)
    elif ext == ".docx":
        return _read_docx(path)
    else:
        raise FileHandlerError(
            f"Unsupported file type: '{ext}'. "
            f"Supported: {', '.join(sorted(SUPPORTED_READ_EXTENSIONS))}"
        )


def _read_text(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        # Fallback for files with non-UTF-8 encoding
        return path.read_text(encoding="latin-1")


def _read_docx(path: Path) -> str:
    try:
        doc = Document(str(path))
        paragraphs = [para.text for para in doc.paragraphs]
        return "\n\n".join(p for p in paragraphs if p.strip())
    except Exception as exc:
        raise FileHandlerError(f"Cannot read .docx file: {exc}") from exc


# ------------------------------------------------------------------
# Writing
# ------------------------------------------------------------------

def write_markdown(text: str, path: Path) -> None:
    """Write text to a .md file.

    Args:
        text: Markdown-formatted string.
        path: Destination file path.
    """
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8")
    logger.info("Wrote markdown: %s", path)


def write_docx(text: str, path: Path) -> None:
    """Write plain or markdown text to a .docx file.

    Markdown headings (# / ## / ###) are mapped to Word heading styles.
    All other lines are written as Normal paragraphs.

    Args:
        text: Plain or markdown-formatted string.
        path: Destination .docx file path.
    """
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    for line in text.splitlines():
        stripped = line.rstrip()
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        else:
            doc.add_paragraph(stripped)

    doc.save(str(path))
    logger.info("Wrote docx: %s", path)


# ------------------------------------------------------------------
# Conversion helpers
# ------------------------------------------------------------------

def markdown_to_html(text: str) -> str:
    """Convert markdown string to HTML."""
    return markdown2.markdown(text, extras=["tables", "fenced-code-blocks"])


def html_to_markdown(html: str) -> str:
    """Convert HTML string to markdown."""
    return md_from_html(html)
